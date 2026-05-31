"""
논문 설계도 .docx 수정 스크립트 (Phase 0)
CACTUS 신규 제안 → H-CACTUS 하이브리드 알고리즘 표현 정정 + 인용 추가
"""
import shutil
from docx import Document
from docx.oxml import OxmlElement

SOURCE = "/Users/kkh/Desktop/2.kkh/2.Study/2.AI연습/1.ppt_team_agent/output/20260518_김경호(2024713673)_아웃라인(논문_설계도)_수정보완.docx"
BACKUP = "/Users/kkh/Desktop/2.kkh/2.Study/2.AI연습/1.ppt_team_agent/output/20260518_김경호(2024713673)_아웃라인(논문_설계도)_수정보완_원본백업_2026-05-20.docx"


def replace_in_para(para, old: str, new: str) -> bool:
    """단락 내 텍스트 교체 (가능하면 run 포맷 유지)"""
    full = para.text
    if old not in full:
        return False
    # 단일 run 안에 있으면 해당 run만 교체
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # 여러 run에 걸친 경우: 첫 run에 전체 교체 텍스트, 나머지 run 비움
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for run in para.runs[1:]:
            run.text = ""
    return True


def insert_paragraph_after(ref_para, text: str) -> None:
    """ref_para 바로 뒤에 새 단락 삽입"""
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    new_p.append(new_r)
    ref_para._element.addnext(new_p)


def main():
    # 1. 백업
    shutil.copy2(SOURCE, BACKUP)
    print(f"백업 생성: {BACKUP}")

    doc = Document(SOURCE)
    paras = doc.paragraphs

    # 2. 단락별 텍스트 교체 매핑
    CHANGES = {
        4: [
            (
                "CACTUS 다중 에이전트 경로 최적화를 결합한",
                "H-CACTUS(Hierarchical CACTUS) 다중 에이전트 경로 최적화를 결합한",
            )
        ],
        9: [
            (
                "다중 AMR 협력 경로 최적화 알고리즘 CACTUS를 통합한",
                "기존 CACTUS(Phan, AAMAS 2024)를 확장한 H-CACTUS 하이브리드 알고리즘"
                "(MILP 작업 배분 + CACTUS 분산 정책 + CBS 충돌 복구 3계층 결합)을 통합한",
            )
        ],
        11: [
            (
                "CACTUS는 동일 SimPy 환경에서 A* Baseline 대비 평균 반송 시간을 20% 이상 단축하고, LLM 기반 노코드 룰 빌더는",
                "H-CACTUS는 동일 SimPy 환경에서 A* Baseline 대비 평균 반송 시간을 20% 이상 단축, "
                "CACTUS·CBS-TS 단독 사용 대비 성공률 향상이 예상된다. LLM 기반 노코드 룰 빌더는",
            )
        ],
        12: [
            (
                "직접 학습하는 CTDE 기반 MAPF 알고리즘과 LLM 자연어 인터페이스를 결합한",
                "구성되는 MILP+MARL+CBS 3중 결합 하이브리드 알고리즘(H-CACTUS)과 LLM 자연어 인터페이스를 결합한",
            )
        ],
        23: [
            (
                "(C3) 4종 알고리즘 비교 프레임워크(CACTUS 신규 제안)",
                "(C3) 4종 baseline + H-CACTUS 신규 하이브리드 알고리즘 제안"
                "(CACTUS Phan AAMAS 2024 및 CBS-TS arXiv 2510.21738 계층 결합)",
            )
        ],
        37: [
            (
                "실 공장 방향 가중 그래프에서 CTDE 기반 협력 학습을 수행하는 CACTUS의 학술적 차별성을 제시한다.",
                "CBS-TS·CACTUS 등 기존 접근의 한계를 정리하고, 이들을 계층적으로 결합한 "
                "H-CACTUS의 학술적 차별성(MILP+MARL+CBS 3중 통합)을 제시한다.",
            )
        ],
        38: [
            (
                "CBS-TS, PPO 단일 에이전트 RL, QMIX 계열 MARL 선행 연구를 정리하고 실 공장 비정형 토폴로지에서의 한계를 분석한다.",
                "CBS-TS(arXiv 2510.21738), CACTUS(Phan AAMAS 2024), PPO, QMIX 계열 MARL 선행 연구를 정리하고 "
                "실 공장 비정형 토폴로지에서의 한계 및 하이브리드 결합 가능성을 분석한다.",
            )
        ],
        54: [
            (
                "A*(Baseline), PPO, CBS-TS, CACTUS를 동일한 Strategy Registry 인터페이스로 통합하여 공정한 비교 실험 환경을 구성하고, "
                "CACTUS가 실 공장 방향 가중 그래프에서 기존 알고리즘 대비 우위를 가짐을 보인다.",
                "A*(Baseline), PPO, CBS-TS, CACTUS, H-CACTUS를 동일한 Strategy Registry 인터페이스로 통합하여 "
                "공정한 비교 실험 환경을 구성하고, 본 연구 제안 알고리즘 H-CACTUS가 실 공장 방향 가중 그래프에서 "
                "기존 4종 알고리즘 대비 성공률·계산시간 양면 우위를 가짐을 보인다.",
            )
        ],
        55: [
            (
                "[Table 2] 4종 알고리즘 이론 비교표  /  [Figure 4] Strategy Registry 프레임워크",
                "[Table 2] 5종 알고리즘 이론 비교표 (A*/PPO/CBS-TS/CACTUS/H-CACTUS)  /  [Figure 4] Strategy Registry 프레임워크",
            )
        ],
        59: [
            (
                "(CACTUS) 제안 알고리즘 — CTDE 기반 MAPF",
                "(CACTUS) Baseline 4 — Phan et al., AAMAS 2024. "
                "Confidence-based Auto-Curriculum for Team Update Stability. CTDE 기반 MAPF",
            )
        ],
        60: [
            (
                "Cooperative Adaptive Control via Training with Unified Scheduling. "
                "QMIX Hypernetwork Mixer + Reverse Curriculum(쉬운 Task → 어려운 Task 단계적 학습). "
                "GraphMAPFEnv로 88노드 방향 가중 그래프 직접 학습.",
                "QMIX Hypernetwork Mixer + Reverse Curriculum(μ_R − η·σ_R ≥ U 충족 시 난이도 상승) 기반 CTDE MAPF. "
                "GraphMAPFEnv로 88노드 방향 가중 그래프 직접 학습.",
            )
        ],
        77: [
            (
                "주장: CACTUS는 A* Baseline 대비 avg_transfer_time 20% 이상 개선을 달성하며, "
                "특히 고밀도(ρ=0.7) 조건에서 혼잡 회피 효과가 두드러짐을 정량 데이터로 보인다.",
                "주장: H-CACTUS는 A* Baseline 대비 avg_transfer_time 20% 이상 개선 및 CACTUS 단독 대비 성공률 향상을 달성하며, "
                "특히 고밀도(ρ=0.7) 조건에서 Conflict Repair 효과가 두드러짐을 정량 데이터로 보인다.",
            )
        ],
        78: [
            (
                "▶ [Figure 7] CACTUS 학습 곡선 — 1차(μ=−4,428) vs. 2차(μ=−427)",
                "▶ [Figure 7] H-CACTUS 컴포넌트별 기여도 — CACTUS 단독 vs. +CBS Repair vs. +Confidence Gate vs. 전체 H-CACTUS",
            )
        ],
        79: [
            (
                "▶ [Figure 8] AMR={8,16,32} × ρ={0.3,0.5,0.7} 조건별 avg_transfer_time",
                "▶ [Figure 8] AMR={8,16,32} × ρ={0.3,0.5,0.7} 조건별 avg_transfer_time — A*/PPO/CBS-TS/CACTUS/H-CACTUS 5종 비교",
            )
        ],
        87: [
            (
                "CACTUS 기반 AI 경로 최적화를 통합한",
                "H-CACTUS 하이브리드 알고리즘 기반 AI 경로 최적화를 통합한",
            )
        ],
        88: [
            (
                "CACTUS 본 학습(Task 028 이후: AMR N=8, bfs_dist 확장, MPS 가속 재시도) 완료 및 CBS-TS+CACTUS 동적 앙상블 전략으로 대규모 에이전트 환경 성능 개선.",
                "H-CACTUS 하이브리드 알고리즘(L1: MILP 작업 배분 + L2: CACTUS 분산 정책 + L3: CBS 충돌 복구 + L4: Confidence Gate) "
                "구현 및 5종 알고리즘 비교 실험 완료. 대규모 에이전트 환경 성능 개선.",
            )
        ],
    }

    # 3. 단락별 교체 실행
    for idx, changes in CHANGES.items():
        para = paras[idx]
        for old, new in changes:
            ok = replace_in_para(para, old, new)
            status = "✓" if ok else "✗ (텍스트 없음)"
            print(f"  [{idx:03d}] {status}: {old[:50]}...")

    # 4. H-CACTUS 신규 단락 삽입 (Figure 5 단락 이후)
    # para[61] = "▶ [Figure 5] CACTUS GraphMAPFEnv 구조 + QMIX Hypernetwork Mixer"
    fig5_para = paras[61]
    h_cactus_para1 = (
        "(H-CACTUS) 본 연구 제안 하이브리드 알고리즘 — CACTUS(L2: 분산 정책) 위에 CBS-TS의 MILP Task Scheduling(L1)과 "
        "Localized CBS Conflict Repair(L3) + Confidence Gate(τ=0.6, L4)를 계층적으로 결합. "
        "선행 LNS2+RL/PRIMAL/EPH 대비 MILP+MARL+CBS 3중 결합은 본 연구 최초 보고."
    )
    h_cactus_para2 = "▶ [Figure 5-bis] H-CACTUS 4계층 아키텍처 (MILP→CACTUS 분산정책→CBS Conflict Repair→A* Fallback)"

    insert_paragraph_after(fig5_para, h_cactus_para2)
    insert_paragraph_after(fig5_para, h_cactus_para1)
    print("  [061+] H-CACTUS 신규 단락 2개 삽입 완료")

    # 5. 참고문헌 추가 (문서 끝에)
    last_para = doc.paragraphs[-1]
    new_refs = [
        "[12] Phan, T. (2024). Confidence-Based Curriculum Learning for Multi-Agent Path Finding. AAMAS 2024. arXiv:2401.05860.",
        "[13] (CBS-TS) Collaborative Task Assignment, Sequencing and Multi-agent Path-finding for Heterogeneous Robots. arXiv:2510.21738.",
        "[14] Wang, Y., et al. (2024). LNS2+RL: Combining Large Neighborhood Search and Learning for Multi-Agent Path Finding. AAAI 2025.",
    ]
    for ref in new_refs:
        insert_paragraph_after(last_para, ref)
        last_para = doc.paragraphs[-1]
        print(f"  참고문헌 추가: {ref[:60]}...")

    # 6. 저장
    doc.save(SOURCE)
    print(f"\n저장 완료: {SOURCE}")


if __name__ == "__main__":
    main()
